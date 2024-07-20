import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi
from googletrans import Translator
from docx import Document
from htmldocx import HtmlToDocx
from io import BytesIO
import markdown2
import re
from lib.constants import price_info

def translate_text(text, dest_lang):
    translator = Translator()
    detected = translator.detect(text)
    if detected == dest_lang:
        return text
    translation = translator.translate(text, dest_lang)
    return translation.text

def extract_transcript(video_id):
    transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['ko', 'en'])
    transcript_text = " ".join([entry['text'] for entry in transcript])
    transcript_text = translate_text(transcript_text, "en")
    return transcript_text

def print_messages():
    if "messages" in st.session_state and len(st.session_state["messages"]) > 0:
        for chat_message in st.session_state["messages"]:
            st.chat_message(chat_message.role).write(chat_message.content)

def stream_parser_default(stream):
    for chunk in stream:
        yield chunk.content

def init_session():
    if "messages" in st.session_state:
        del st.session_state.messages
    if "chat_history" in st.session_state:
        del st.session_state.chat_history
    if "llm" in st.session_state:
        del st.session_state.llm
    if "transcript" in st.session_state:
        del st.session_state.transcript
    if "video_id" in st.session_state:
        del st.session_state.video_id

# Function to save chat history to a docx file
def save_chat_to_docx(chat_history):
    doc = Document()
    doc.add_heading("Chat History", 0)
    
    h2d = HtmlToDocx()

    # Iterate over chat history and format it
    for speaker, chat in chat_history[1:]:  # Skip the first two entries
        html_chat = markdown2.markdown(chat)
        if speaker == "human":
            p = doc.add_paragraph()
            p.add_run("You: ").bold = True
            h2d.add_html_to_document(html_chat, doc)
        else:
            p = doc.add_paragraph()
            p.add_run("AI: ").bold = True
            h2d.add_html_to_document(html_chat, doc)
    
    # Save the docx file in memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def update_token_usage(usage):
    st.session_state.total_tokens += usage['total_tokens']
    st.session_state.prompt_tokens += usage['prompt_tokens']
    st.session_state.completion_tokens += usage['completion_tokens']

def show_cost(model):
    input_token_price = price_info[model]["input"]
    output_token_price = price_info[model]["output"]
    input_cost = st.session_state.prompt_tokens * input_token_price
    output_cost = st.session_state.completion_tokens * output_token_price
    total_cost = input_cost + output_cost
    st.sidebar.markdown(f"""
### Cost Breakdown
**Model:** {model}

**Input Tokens Used:** {st.session_state.prompt_tokens}  
**Input Token Cost (per 1M tokens):** \${input_token_price * 1_000_000:.2f}  
**Input Cost:** ${input_cost:.4f}

**Output Tokens Used:** {st.session_state.completion_tokens}  
**Output Token Cost (per 1M tokens):** \${output_token_price * 1_000_000:.2f}  
**Output Cost:** ${output_cost:.4f}

**Total Cost (USD):** ${total_cost:.4f}
""")
    
def extract_video_id(url):
    if is_valid_url(url):
        
        pattern = r"(?:v=|\/)([0-9A-Za-z_-]{11}).*"
        match = re.search(pattern, url)
        if match:
            return match.group(1)
        return None
    else:
        return url

def is_valid_url(url):
    pattern = re.compile(
        r'^(?:http|ftp)s?://' # http:// or https://
        r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+(?:[A-Z]{2,6}\.?|[A-Z0-9-]{2,}\.?)|' # domain...
        r'localhost|' # localhost...
        r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}|' # ... IPv4 address
        r'\[?[A-F0-9]*:[A-F0-9:]+\]?)' # ... IPv6 address
        r'(?::\d+)?' # port
        r'(?:/?|[/?]\S+)$', re.IGNORECASE)
    return re.match(pattern, url) is not None